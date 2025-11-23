"""
Servicio de OCR para extracción de texto de PDFs y documentos Word
Soporta PDFs nativos digitales, escaneados, y documentos Word (.docx)
"""
import io
import fitz  # PyMuPDF
from typing import Optional
import numpy as np
from PIL import Image

# Fix para compatibilidad con Pillow 10+ (ANTIALIAS fue removido)
try:
    from PIL.Image import Resampling
    # Crear alias para compatibilidad con código antiguo (EasyOCR puede usar ANTIALIAS)
    Image.ANTIALIAS = Resampling.LANCZOS
except (ImportError, AttributeError):
    # Si Resampling no existe, usar LANCZOS directamente
    try:
        Image.ANTIALIAS = Image.LANCZOS
    except AttributeError:
        pass

import easyocr
from docx import Document  # python-docx para .docx


class OCRService:
    """Servicio para extracción de texto de documentos PDF"""
    
    def __init__(self):
        # Inicializar EasyOCR solo cuando sea necesario (para PDFs escaneados)
        self.easyocr_reader = None
        self.debug_logs = []  # Logs de depuración
    
    def _add_log(self, message, level="INFO"):
        """Añade un log a la lista de logs de depuración"""
        log_entry = f"[{level}] {message}"
        print(log_entry)
        self.debug_logs.append(log_entry)
    
    def get_logs(self):
        """Obtiene los logs de depuración y los limpia"""
        logs = self.debug_logs.copy()
        self.debug_logs = []
        return logs
    
    async def extract_text(self, file_content: bytes, filename: Optional[str] = None) -> str:
        """
        Extrae texto de un documento (PDF, DOC, DOCX)
        
        Args:
            file_content: Contenido del archivo en bytes
            filename: Nombre del archivo (opcional)
        
        Returns:
            Texto extraído del documento
        """
        try:
            # Detectar tipo de archivo por extensión
            if filename:
                filename_lower = filename.lower()
                if filename_lower.endswith('.docx'):
                    return await self._extract_from_docx(file_content)
                elif filename_lower.endswith('.doc'):
                    # .doc antiguo - python-docx NO puede leer .doc, solo .docx
                    # Intentar de todas formas por si acaso, pero probablemente fallará
                    try:
                        return await self._extract_from_docx(file_content)
                    except Exception as e:
                        error_msg = str(e).lower()
                        if "not a zip file" in error_msg or "bad zipfile" in error_msg:
                            raise Exception(
                                "Los archivos .doc (formato antiguo de Word) no están soportados directamente. "
                                "Por favor, convierte el archivo a .docx o PDF:\n"
                                "1. Abre el archivo en Microsoft Word\n"
                                "2. Ve a 'Archivo' → 'Guardar como'\n"
                                "3. Selecciona formato '.docx' o 'PDF'\n"
                                "4. Guarda y vuelve a subir el archivo convertido"
                            )
                        raise
            
            # Por defecto, tratar como PDF
            try:
                # Intentar extracción directa con PyMuPDF (PDFs nativos)
                text = await self._extract_with_pymupdf(file_content)
                content_length = len(text.strip())
                text_lower = text.lower()
                
                self._add_log(f"Texto extraído con PyMuPDF: {content_length} caracteres")
                self._add_log(f"Primeros 200 caracteres: {text[:200]}")
                
                # Verificar si el texto extraído es solo metadatos/encabezado
                header_keywords = ["copia autentica", "localizador", "registro salida", "fecha registro", "sello", "acceda a la página", "acceda a la pagina", "para visualizar el documento"]
                has_header = any(keyword in text_lower for keyword in header_keywords)
                
                # Verificar si hay enlaces a documentos externos (indica que el contenido real está en otra URL)
                has_external_link = "verdocumentos" in text_lower or "visualizar el documento" in text_lower or "jcyl.es" in text_lower
                
                # Palabras clave que indican contenido real del documento (no solo metadatos)
                # Excluir "discapacidad" si solo aparece en contexto de registro/trámite
                content_keywords = ["resolución", "determina que", "diagnóstico", "m75", "lesión", "hombro", "anexo", "baremo", "grado de discapacidad", "reconocimiento del grado"]
                has_content = False
                for keyword in content_keywords:
                    if keyword in text_lower:
                        # Verificar que no esté en contexto de registro/trámite
                        keyword_index = text_lower.find(keyword)
                        if keyword_index >= 0:
                            # Buscar contexto alrededor
                            context_start = max(0, keyword_index - 100)
                            context_end = min(len(text_lower), keyword_index + len(keyword) + 100)
                            context = text_lower[context_start:context_end]
                            # Si el contexto no es principalmente metadatos, considerar contenido real
                            if not all(meta_word in context for meta_word in ["registro", "localizador", "fecha registro", "sello"]):
                                has_content = True
                                break
                
                self._add_log(f"Tiene encabezado: {has_header}, Tiene contenido real: {has_content}, Tiene enlace externo: {has_external_link}")
                
                # Si tiene enlace externo, es probable que el contenido real esté en otra URL
                if has_external_link:
                    self._add_log(f"⚠️ ATENCIÓN: Este PDF parece ser una 'copia auténtica' con enlace externo.", "WARNING")
                    self._add_log(f"El contenido real del documento puede estar en una URL externa.", "WARNING")
                    # Intentar extraer la URL
                    import re
                    url_pattern = r'https?://[^\s]+'
                    urls = re.findall(url_pattern, text)
                    if urls:
                        self._add_log(f"URL encontrada: {urls[0]}", "WARNING")
                        self._add_log(f"El documento real puede estar en: {urls[0]}", "WARNING")
                
                # Si tiene encabezado pero NO tiene contenido real, FORZAR OCR
                # O si el texto es muy corto (< 1000 caracteres)
                # O si el texto tiene más de 1000 caracteres pero solo contiene metadatos (más del 80% es encabezado)
                should_use_ocr = False
                
                # Calcular porcentaje de texto que es encabezado
                header_text_length = 0
                for keyword in header_keywords:
                    # Contar ocurrencias y longitud aproximada
                    count = text_lower.count(keyword)
                    header_text_length += count * len(keyword) * 2  # Aproximación
                
                header_percentage = (header_text_length / content_length * 100) if content_length > 0 else 0
                
                # Si tiene enlace externo y no tiene contenido real, es muy probable que necesite OCR
                if has_external_link and not has_content:
                    self._add_log(f"PDF con enlace externo sin contenido real. FORZANDO OCR para intentar extraer contenido...", "WARNING")
                    should_use_ocr = True
                elif has_header and not has_content:
                    self._add_log(f"Solo se detectó encabezado sin contenido real. FORZANDO OCR...", "WARNING")
                    should_use_ocr = True
                elif content_length < 1000:
                    self._add_log(f"Texto muy corto ({content_length} caracteres). Intentando OCR...", "WARNING")
                    should_use_ocr = True
                elif has_header and header_percentage > 80:
                    self._add_log(f"Texto parece ser principalmente metadatos ({header_percentage:.1f}% encabezado). FORZANDO OCR...", "WARNING")
                    should_use_ocr = True
                
                if should_use_ocr:
                    try:
                        self._add_log(f"Iniciando extracción con OCR...", "WARNING")
                        self._add_log(f"Esto puede tardar varios minutos la primera vez (carga del modelo)...", "WARNING")
                        ocr_text = await self._extract_with_ocr(file_content)
                        ocr_length = len(ocr_text.strip())
                        self._add_log(f"OCR extrajo {ocr_length} caracteres vs {content_length} con PyMuPDF")
                        
                        # Mostrar muestra del texto OCR
                        if ocr_length > 0:
                            self._add_log(f"Primeros 300 caracteres del OCR: {ocr_text[:300]}")
                        
                        # Si OCR extrajo más texto o encontró contenido real, usarlo
                        ocr_lower = ocr_text.lower()
                        ocr_has_content = any(keyword in ocr_lower for keyword in content_keywords)
                        
                        self._add_log(f"OCR tiene contenido real: {ocr_has_content}")
                        
                        if ocr_length > content_length or (ocr_has_content and not has_content):
                            self._add_log(f"Usando texto de OCR (mejor contenido detectado)", "SUCCESS")
                            return ocr_text
                        else:
                            self._add_log(f"OCR no mejoró la extracción, usando texto original", "WARNING")
                            # Aun así, si OCR extrajo algo, intentar combinarlo
                            if ocr_length > 100:
                                self._add_log(f"Combinando texto PyMuPDF + OCR...")
                                return text + "\n\n" + ocr_text
                    except Exception as ocr_error:
                        self._add_log(f"Error en OCR: {str(ocr_error)}", "ERROR")
                        import traceback
                        error_trace = traceback.format_exc()
                        self._add_log(f"Traceback completo: {error_trace}", "ERROR")
                        # Continuar con el texto extraído aunque sea poco
                
                return text
            except Exception as pdf_error:
                # Si falla como PDF, puede ser que el archivo esté corrupto o no sea PDF
                raise Exception(f"Error procesando como PDF: {str(pdf_error)}. Verifica que el archivo sea válido.")
        
        except Exception as e:
            error_msg = str(e)
            # Mejorar mensajes de error
            if "fitz" in error_msg.lower() or "pymupdf" in error_msg.lower():
                raise Exception("Error al leer el PDF. El archivo puede estar corrupto o protegido.")
            elif "docx" in error_msg.lower():
                raise Exception("Error al leer el documento Word. Verifica que sea un archivo .docx válido.")
            else:
                raise Exception(f"Error en extracción de texto: {error_msg}")
    
    async def _extract_with_pymupdf(self, pdf_content: bytes) -> str:
        """Extrae texto de PDFs nativos digitales usando PyMuPDF"""
        doc = fitz.open(stream=pdf_content, filetype="pdf")
        text_parts = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Método 1: Extracción estándar
            text = page.get_text()
            
            # Método 2: Si no hay suficiente texto, intentar con diferentes opciones
            if len(text.strip()) < 100:
                # Intentar con opciones de extracción más agresivas
                text_dict = page.get_text("dict")
                # Extraer texto de bloques
                blocks_text = []
                for block in text_dict.get("blocks", []):
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line.get("spans", []):
                                if span.get("text"):
                                    blocks_text.append(span["text"])
                if blocks_text:
                    text = " ".join(blocks_text)
            
            # Método 3: Intentar extraer texto de anotaciones
            if len(text.strip()) < 100:
                annotations_text = []
                for annot in page.annots():
                    if annot.type[0] == 2:  # Text annotation
                        info = annot.info
                        if info.get("content"):
                            annotations_text.append(info["content"])
                if annotations_text:
                    text = text + "\n" + "\n".join(annotations_text)
            
            # Método 4: Extraer texto con diferentes opciones de layout
            if len(text.strip()) < 100:
                # Intentar extracción con diferentes flags
                text_flags = page.get_text("text", flags=11)  # flags para mejor extracción
                if len(text_flags.strip()) > len(text.strip()):
                    text = text_flags
            
            # Método 5: Extraer texto de formularios/widgets
            if len(text.strip()) < 100:
                widgets_text = []
                for widget in page.widgets():
                    if widget.field_value:
                        widgets_text.append(str(widget.field_value))
                if widgets_text:
                    text = text + "\n" + "\n".join(widgets_text)
            
            text_parts.append(text)
        
        doc.close()
        full_text = "\n\n".join(text_parts)
        
        # Si aún no hay suficiente texto, puede ser que el PDF tenga el contenido en imágenes
        # En ese caso, retornar lo que tenemos pero marcar que necesita OCR
        return full_text
    
    async def _extract_with_ocr(self, pdf_content: bytes) -> str:
        """Extrae texto de PDFs escaneados usando EasyOCR"""
        self._add_log(f"Inicializando EasyOCR...")
        try:
            if self.easyocr_reader is None:
                # Inicializar EasyOCR con español e inglés
                self._add_log(f"Cargando modelo EasyOCR (esto puede tardar varios minutos la primera vez)...", "WARNING")
                self._add_log(f"Por favor, espere...", "WARNING")
                try:
                    self.easyocr_reader = easyocr.Reader(['es', 'en'], gpu=False)
                    self._add_log(f"Modelo EasyOCR cargado correctamente", "SUCCESS")
                except Exception as init_error:
                    self._add_log(f"Error al cargar EasyOCR: {str(init_error)}", "ERROR")
                    raise Exception(f"No se pudo inicializar EasyOCR. Verifica que esté instalado correctamente: {str(init_error)}")
        except Exception as e:
            self._add_log(f"Error crítico en inicialización de OCR: {str(e)}", "ERROR")
            raise
        
        doc = fitz.open(stream=pdf_content, filetype="pdf")
        text_parts = []
        total_pages = len(doc)
        
        self._add_log(f"Procesando {total_pages} página(s) con OCR...")
        
        for page_num in range(total_pages):
            page = doc[page_num]
            self._add_log(f"Procesando página {page_num + 1}/{total_pages}...")
            
            # Convertir página a imagen con mayor resolución para mejor OCR
            # Matrix(3, 3) = 3x zoom para mejor calidad
            pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            img_array = np.array(img)
            
            # Realizar OCR
            self._add_log(f"Ejecutando OCR en página {page_num + 1}...")
            results = self.easyocr_reader.readtext(img_array)
            page_text = " ".join([result[1] for result in results])
            self._add_log(f"Página {page_num + 1}: {len(page_text)} caracteres extraídos")
            text_parts.append(page_text)
        
        doc.close()
        full_text = "\n\n".join(text_parts)
        self._add_log(f"Total extraído: {len(full_text)} caracteres", "SUCCESS")
        return full_text
    
    async def _extract_from_docx(self, docx_content: bytes) -> str:
        """
        Extrae texto de un archivo .docx
        
        Args:
            docx_content: Contenido del archivo .docx en bytes
        
        Returns:
            Texto extraído del documento
        """
        try:
            # Verificar que el contenido no esté vacío
            if not docx_content or len(docx_content) == 0:
                raise Exception("El archivo está vacío")
            
            # Crear un objeto de archivo en memoria
            docx_file = io.BytesIO(docx_content)
            
            # Verificar que el archivo sea realmente un .docx (debe empezar con PK, que es la firma ZIP)
            if len(docx_content) < 4 or docx_content[:2] != b'PK':
                raise Exception(
                    "El archivo no es un documento Word válido (.docx). "
                    "Los archivos .doc (formato antiguo de Word 97-2003) no están soportados. "
                    "Por favor, convierte el archivo:\n"
                    "1. Abre el archivo en Microsoft Word\n"
                    "2. Ve a 'Archivo' → 'Guardar como'\n"
                    "3. Selecciona formato '.docx' o 'PDF'\n"
                    "4. Guarda y vuelve a subir el archivo convertido"
                )
            
            # Abrir el documento Word
            try:
                doc = Document(docx_file)
            except Exception as e:
                error_str = str(e).lower()
                if "not a zip file" in error_str or "bad zipfile" in error_str or "invalid" in error_str:
                    # Probablemente es un .doc antiguo o archivo corrupto
                    raise Exception(
                        "El archivo no es un documento Word válido (.docx). "
                        "Si es un archivo .doc (formato antiguo), necesitas convertirlo a .docx o PDF. "
                        "Si es un .docx, el archivo puede estar corrupto. "
                        "Intenta abrirlo en Word y guardarlo de nuevo."
                    )
                raise Exception(f"Error al abrir el documento Word: {str(e)}")
            
            # Extraer texto de todos los párrafos
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            extracted_text = "\n\n".join(text_parts)
            
            if not extracted_text or len(extracted_text.strip()) < 1:
                raise Exception("El documento Word no contiene texto extraíble")
            
            return extracted_text
        
        except Exception as e:
            error_msg = str(e)
            if "python-docx" in error_msg.lower() or "document" in error_msg.lower():
                raise Exception(f"Error extrayendo texto de documento Word: {error_msg}. Verifica que el archivo sea un .docx válido.")
            raise
    
    async def detect_pdf_type(self, pdf_content: bytes) -> str:
        """
        Detecta si el PDF es nativo digital o escaneado
        
        Returns:
            'native' o 'scanned'
        """
        try:
            doc = fitz.open(stream=pdf_content, filetype="pdf")
            page = doc[0]
            text = page.get_text()
            doc.close()
            
            # Si hay texto suficiente, es nativo
            if len(text.strip()) > 50:
                return "native"
            else:
                return "scanned"
        
        except Exception:
            return "scanned"

